"""
informes.py — Lógica de minutas e informes DOCX para la PWA
"""
import json
import os
import re
from datetime import datetime, timedelta, timezone
from pathlib import Path

_DATA        = Path(os.environ.get("DATA_DIR", Path(__file__).parent.resolve()))
MINUTAS_FILE = _DATA / "minutas.json"
VENEZUELA_TZ = timezone(timedelta(hours=-4))

OBSERVACION_DEFAULT = (
    "La información fue notificada a la digna superioridad en tiempo real "
    "para su conocimiento, evaluación y fines consiguientes."
)

INTRO_INFORMACION_INICIAL = (
    "Mediante monitoreo a través del patrullaje cibernético, por diversos portales de "
    "información web y redes sociales, a los fines de identificar informaciones de interés "
    "operativo, relacionado al campo de la inteligencia y contrainteligencia, con el fin de "
    "informar para recomendar acciones o su efecto informar para neutralizar alguna célula u "
    "organización que pretendan integrar alguna inteligencia regional, que busque sustraer y "
    "suministrar información en contra del gobierno bolivariano, las fuerzas de seguridad "
    "venezolanas y los sectores populares, todo esto con el objetivo de causar malestar en la "
    "población, para volver a generar un clima de protestas."
)

EDS_MARACAIBO = [
    "E/S SERVICIOS POPULARES",
    "E/S AUTOMOTRIZ",
    "E/S MILAGROS",
    "E/S CALZADA",
    "E/S PICHINCHA",
    "E/S NIGALE",
    "E/S CARMEN",
    "E/S DELICIAS",
]

ANALISIS_TEXTOS = [
    "La información difundida en redes sociales representa un monitoreo continuo de eventos relevantes en el estado Zulia, permitiendo evaluar dinámicas sociales, políticas e institucionales.",
    "Este reporte contribuye a la comprensión del panorama informativo y el impacto de las acciones institucionales en la opinión pública.",
    "El análisis de contenidos en plataformas digitales permite identificar tendencias, evaluar alcance de mensajes y respuestas ciudadanas.",
]

OBSERVACION_TEXTOS = [
    "Se recomienda continuar con el monitoreo permanente de redes sociales y fuentes informativas para mantener actualizado el panorama situacional del estado.",
    "Es importante mantener vigilancia sobre la evolución de estos eventos y su potencial impacto en la dinámica social de la región.",
]

RENGLONES_SITUACION = [
    "POLÍTICO",
    "MONITOREO",
    "DEPORTE",
    "ECONÓMICO",
    "AMBIENTAL",
    "SUCESOS",
]

MESES = ["ENERO","FEBRERO","MARZO","ABRIL","MAYO","JUNIO",
         "JULIO","AGOSTO","SEPTIEMBRE","OCTUBRE","NOVIEMBRE","DICIEMBRE"]

MESES_TITULO = ["Enero","Febrero","Marzo","Abril","Mayo","Junio",
                "Julio","Agosto","Septiembre","Octubre","Noviembre","Diciembre"]


# ── Persistencia de minutas ──────────────────────────────────────────────────

def cargar_minutas():
    try:
        if MINUTAS_FILE.is_file():
            with MINUTAS_FILE.open(encoding="utf-8") as f:
                data = json.load(f)
                return data if isinstance(data, list) else []
    except Exception:
        pass
    return []


def guardar_minutas(minutas):
    MINUTAS_FILE.parent.mkdir(parents=True, exist_ok=True)
    with MINUTAS_FILE.open("w", encoding="utf-8") as f:
        json.dump(minutas, f, ensure_ascii=False, indent=2)


# ── Fechas (hora Venezuela UTC-4) ────────────────────────────────────────────

def obtener_fecha_formato(fecha=None):
    if fecha is None:
        fecha = datetime.now(VENEZUELA_TZ)
    return f"{fecha.day} DE {MESES[fecha.month - 1]} {fecha.year}"


def obtener_fecha_titulo(fecha=None):
    """Formato: 03 De Junio año 2026"""
    if fecha is None:
        fecha = datetime.now(VENEZUELA_TZ)
    return f"{str(fecha.day).zfill(2)} De {MESES_TITULO[fecha.month - 1]} año {fecha.year}"


def obtener_rango_horario():
    ahora = datetime.now(VENEZUELA_TZ)
    h = ahora.hour
    if 8 <= h < 17:
        f = obtener_fecha_formato(ahora)
        return f"{f} 08:00HRS – {f} 17:00HRS", "diurno"
    elif h >= 17:
        f1 = obtener_fecha_formato(ahora)
        f2 = obtener_fecha_formato(ahora + timedelta(days=1))
        return f"{f1} 17:00HRS – {f2} 08:00HRS", "nocturno"
    else:
        f1 = obtener_fecha_formato(ahora - timedelta(days=1))
        f2 = obtener_fecha_formato(ahora)
        return f"{f1} 17:00HRS – {f2} 08:00HRS", "nocturno"


def nombre_archivo_informe():
    hoy = datetime.now(VENEZUELA_TZ)
    return f"INFORME_{hoy.strftime('%H')}00HRS_ENTORNO_ZULIA_{str(hoy.day).zfill(2)}{str(hoy.month).zfill(2)}{hoy.year}.docx"


def nombre_archivo_situacion():
    hoy = datetime.now(VENEZUELA_TZ)
    return f"INFORME_SITUACION_OPERATIVA_{str(hoy.day).zfill(2)}{str(hoy.month).zfill(2)}{hoy.year}.docx"


# ── Helpers DOCX ─────────────────────────────────────────────────────────────

def _header_oscuro(doc, texto, color="1F2D6E"):
    """Párrafo con fondo oscuro y texto blanco."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    pPr.append(shd)
    run = p.add_run(texto)
    run.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.size = Pt(11)
    return p


def _limpiar_hecho(texto):
    """Elimina URLs y prefijos como '1- se informó a la superioridad'."""
    import re

    # Remover URLs
    texto = re.sub(r'https?://\S+|www\.\S+', '', texto)

    # Remover prefijos numéricos como "1- ", "2- ", etc. al inicio o después de espacios
    # También remover variaciones de "se informó a la superioridad"
    texto = re.sub(r'^\d+\-\s*', '', texto)
    texto = re.sub(r'\n\d+\-\s*', '\n', texto)
    texto = re.sub(r'se\s+informó\s+a\s+la\s+superioridad[^.]*\.?\s*', '', texto, flags=re.IGNORECASE)

    # Limpiar espacios en blanco excesivos
    texto = re.sub(r'\s+', ' ', texto).strip()

    return texto


def _agregar_minuta(doc, m):
    """Formato simple para INFORME ESPECIAL: FECHA, HORA, HECHO, ANALISIS, OBSERVACION, 1 foto."""
    from docx.shared import Inches

    doc.add_paragraph(f"FECHA: {m.get('fecha', '')}")
    doc.add_paragraph(f"HORA: {m.get('hora', '')}")
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("HECHO").bold = True
    hecho_limpio = _limpiar_hecho(m.get("hecho", ""))
    doc.add_paragraph(hecho_limpio)

    analisis = m.get("analisis") or "\n".join(ANALISIS_TEXTOS)
    p = doc.add_paragraph()
    p.add_run("ANALISIS:").bold = True
    for linea in analisis.split("\n"):
        if linea.strip():
            doc.add_paragraph(linea.strip())

    obs = m.get("observacion") or "\n".join(OBSERVACION_TEXTOS)
    p = doc.add_paragraph()
    p.add_run("OBSERVACION:").bold = True
    for linea in obs.split("\n"):
        if linea.strip():
            doc.add_paragraph(linea.strip())

    # Solo la PRIMERA foto disponible (nunca URLs en el documento)
    fotos = [x for x in m.get("media", []) if x.get("tipo") == "foto"]
    if fotos:
        fp = Path(fotos[0].get("path", ""))
        if fp.is_file():
            try:
                doc.add_picture(str(fp), width=Inches(5.5))
            except Exception:
                pass

    doc.add_paragraph()


def _agregar_minuta_completa(doc, m):
    """Formato completo para SITUACIÓN OPERATIVA: incluye LUGAR, FUENTE, INCIDENCIA."""
    from docx.shared import Inches

    hora = m.get("hora", "")
    if hora and "HRS" not in hora.upper():
        hora = hora + " HRS"

    doc.add_paragraph(f"FECHA: {m.get('fecha', '')}")
    doc.add_paragraph(f"HORA: {hora}")

    lugar = m.get("lugar", "Estado Zulia")
    doc.add_paragraph(f"LUGAR: {lugar}.")

    fuente = m.get("fuente", "")
    if fuente:
        doc.add_paragraph(f"FUENTE: {fuente}.")

    incidencia = m.get("incidencia", "")
    if incidencia:
        doc.add_paragraph(f"INCIDENCIA: {incidencia}.")

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("HECHO: ").bold = True
    hecho_limpio = _limpiar_hecho(m.get("hecho", ""))
    p.add_run(hecho_limpio)

    doc.add_paragraph()

    analisis = m.get("analisis") or "\n".join(ANALISIS_TEXTOS)
    p = doc.add_paragraph()
    p.add_run("ANÁLISIS: ").bold = True
    p.add_run(analisis)

    doc.add_paragraph()

    obs = m.get("observacion") or OBSERVACION_DEFAULT
    p = doc.add_paragraph()
    p.add_run("OBSERVACIÓN: ").bold = True
    p.add_run(obs)

    # Primera foto disponible
    fotos = [x for x in m.get("media", []) if x.get("tipo") == "foto"]
    if fotos:
        fp = Path(fotos[0].get("path", ""))
        if fp.is_file():
            try:
                doc.add_picture(str(fp), width=Inches(5.5))
            except Exception:
                pass

    doc.add_paragraph()


# ── Informe Especial ─────────────────────────────────────────────────────────

def generar_docx(minutas):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    ahora = datetime.now(VENEZUELA_TZ)

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)

    # CONFIDENCIAL
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CONFIDENCIAL")
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(204, 0, 0)

    # INFORME ESPECIAL
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("INFORME ESPECIAL")
    r.bold = True; r.font.size = Pt(14)

    doc.add_paragraph(f"FECHA: {obtener_fecha_formato(ahora)} {ahora.strftime('%H:%M')}HRS")
    doc.add_paragraph("LUGAR: Estado ZULIA")
    doc.add_paragraph("ASUNTO: Monitoreo")
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("Ámbito Social").bold = True

    for m in minutas:
        _agregar_minuta(doc, m)

    # Electricidad
    p = doc.add_paragraph()
    p.add_run("Electricidad:").bold = True
    doc.add_paragraph(
        "En el estado Zulia, se reporta que la capacidad operativa del sistema eléctrico "
        "se mantiene en un 60%, lo que indica un nivel de funcionamiento estable pero todavía "
        "por debajo de la plena capacidad; esta condición requiere monitoreo constante para "
        "prevenir interrupciones y garantizar la continuidad del suministro tanto a nivel "
        "residencial como industrial, especialmente en sectores críticos que dependen del "
        "servicio eléctrico para su operatividad diaria."
    )

    ruta = _DATA / nombre_archivo_informe()
    doc.save(str(ruta))
    return ruta


# ── Informe de Situación Operativa ───────────────────────────────────────────

def generar_docx_situacion(minutas, asignaciones, homicidios=0, eds_estado=None):
    """
    Genera el Informe de Situación Operativa.
    asignaciones: {str(idx) -> renglón}
    homicidios: entero
    eds_estado: dict {nombre_eds: "FUNCIONANDO"|"NO FUNCIONANDO"} o None para todos FUNCIONANDO
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    ahora  = datetime.now(VENEZUELA_TZ)
    rango, _ = obtener_rango_horario()

    if eds_estado is None:
        eds_estado = {eds: "FUNCIONANDO" for eds in EDS_MARACAIBO}

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)

    # ── Encabezado principal ──────────────────────────────────────────────
    _header_oscuro(doc, "INFORME DE INTELIGENCIA", "000000")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DIVISIÓN DE INTELIGENCIA ESTRATÉGICA ZULIA")
    r.bold = True; r.font.size = Pt(12)

    p = doc.add_paragraph()
    p.add_run("Asunto: ").bold = True
    p.add_run("Secreto")

    p = doc.add_paragraph()
    p.add_run("Fecha: ").bold = True
    p.add_run(obtener_fecha_titulo(ahora))

    doc.add_paragraph()

    # ── INFORMACIÓN INICIAL ───────────────────────────────────────────────
    _header_oscuro(doc, ":     INFORMACIÓN INICIAL")
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(INTRO_INFORMACION_INICIAL)

    doc.add_paragraph()

    # ── BÚSQUEDA Y ────────────────────────────────────────────────────────
    _header_oscuro(doc, "BÚSQUEDA Y")
    doc.add_paragraph()

    # Agrupar minutas por renglón
    grupos = {}
    accidentes = []
    for idx_str, renglon in asignaciones.items():
        if not renglon:
            continue
        try:
            idx = int(idx_str)
            if 0 <= idx < len(minutas):
                if renglon == "ACCIDENTES DE TRÁNSITO":
                    accidentes.append(minutas[idx])
                else:
                    grupos.setdefault(renglon, []).append(minutas[idx])
        except (ValueError, TypeError):
            pass

    # Renglones en orden
    todos_renglon = RENGLONES_SITUACION + [r for r in grupos if r not in RENGLONES_SITUACION]

    for renglon in todos_renglon:
        p = doc.add_paragraph()
        p.add_run(renglon).bold = True

        if renglon in grupos:
            for m in grupos[renglon]:
                _agregar_minuta_completa(doc, m)
        else:
            doc.add_paragraph()

    # ── ABASTECIMIENTO Y COMBUSTIBLE ──────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("ABASTECIMIENTO Y COMBUSTIBLE: ").bold = True
    p.add_run(
        "En cumplimiento de las tareas de monitoreo sobre el funcionamiento de las estaciones "
        "de servicio (E/S) en Maracaibo, se realizó el levantamiento de información actualizada "
        "para el día de hoy."
    )
    doc.add_paragraph()

    for eds, estado in eds_estado.items():
        p = doc.add_paragraph()
        p.add_run(f"{eds}: ").bold = True
        p.add_run(estado)

    doc.add_paragraph()

    # ── ACCIDENTES DE TRÁNSITO ────────────────────────────────────────────
    p = doc.add_paragraph()
    p.add_run("ACCIDENTES DE TRÁNSITO:").bold = True

    if accidentes:
        for m in accidentes:
            _agregar_minuta_completa(doc, m)
    else:
        doc.add_paragraph()

    # ── HOMICIDIOS ────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.add_run("HOMICIDIOS: ").bold = True
    p.add_run(f"({str(homicidios).zfill(2)}).")

    doc.add_paragraph()

    # ── ANEXO ─────────────────────────────────────────────────────────────
    _header_oscuro(doc, "ANEXO", "2E75B6")
    doc.add_paragraph()

    # Fotos de todas las minutas asignadas como evidencia del anexo
    todas_fotos = []
    for idx_str in asignaciones:
        try:
            idx = int(idx_str)
            if 0 <= idx < len(minutas):
                for item in minutas[idx].get("media", []):
                    if item.get("tipo") == "foto":
                        todas_fotos.append(item)
        except (ValueError, TypeError):
            pass

    if todas_fotos:
        p = doc.add_paragraph()
        p.add_run("EVIDENCIA FOTOGRÁFICA:").bold = True
        for foto in todas_fotos:
            fp = Path(foto.get("path", ""))
            if fp.is_file():
                try:
                    doc.add_picture(str(fp), width=Inches(5.5))
                except Exception:
                    pass

    ruta = _DATA / nombre_archivo_situacion()
    doc.save(str(ruta))
    return ruta
